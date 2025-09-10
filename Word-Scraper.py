# Imports
from pathlib import Path
from docx import Document
import pandas as pd

# Inputs
input_path = input("Input DOCX File or Folder Here: ")
input_save = input("Output CSV File as: ")

# Create lists of all documents and paragraphs from input file or folder
document_list = []
paragraph_list = []
file_path = Path(input_path)

if file_path.is_dir():
    for item in file_path.iterdir():
        if item.suffix == '.docx':
            document = Document(item)
            for paragraph in document.paragraphs:
                paragraph_list.append(paragraph.text)
                document_list.append(item.stem)

elif file_path.is_file():
    document = Document(file_path)
    for paragraph in document.paragraphs:
        paragraph_list.append(paragraph.text)
        document_list.append(file_path.stem)

# Save to CSV
data = {"Document":document_list,"Paragraph":paragraph_list}
result = pd.DataFrame(data)
result.to_csv(input_save)
